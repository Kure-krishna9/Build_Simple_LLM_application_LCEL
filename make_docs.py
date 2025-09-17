from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading("FastAPI + LangChain + Groq + LangServe Translator API", level=0)

# Section 1
doc.add_heading("1. Project Overview", level=1)
doc.add_paragraph(
    "This application exposes an API that translates text into a target language using:\n\n"
    "- FastAPI → to create and serve the API.\n"
    "- LangChain → to build a pipeline (prompt + LLM + parser).\n"
    "- Groq → as the backend LLM provider (using Gemma-2-9b-it).\n"
    "- LangServe → to expose LangChain chains as API routes automatically.\n"
    "- dotenv → to securely load the API key."
)

# Section 2: Code Walkthrough
doc.add_heading("2. Code Walkthrough", level=1)

sections = {
    "Imports": """FastAPI, LangChain modules, Groq integration, dotenv for environment variables, and LangServe to expose chains.""",
    "Load API Key": """Loads the Groq API key from .env file.""",
    "Define the Model": """Initializes Groq's Gemma-2-9b-it model using ChatGroq.""",
    "Create Prompt Template": """Defines the system and user messages for translation.""",
    "Output Parser": """Uses StrOutputParser to return clean string outputs.""",
    "Build the Chain": """Connects prompt → model → parser into a pipeline.""",
    "FastAPI App Setup": """Creates FastAPI app with metadata.""",
    "Expose Chain as API": """Adds LangChain runnable chain as a /chain endpoint using add_routes.""",
    "Run the App": """Runs FastAPI with Uvicorn at localhost:8000."""
}

for title, content in sections.items():
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)

# Section 3
doc.add_heading("3. How to Run", level=1)
doc.add_paragraph(
    "1. Install dependencies:\n"
    "   pip install fastapi uvicorn langchain langchain-groq langserve python-dotenv\n\n"
    "2. Create .env file with your Groq API key:\n"
    "   Groq_key=your_groq_api_key_here\n\n"
    "3. Run the server:\n"
    "   python main.py\n\n"
    "Server runs at http://localhost:8000"
)

# Section 4
doc.add_heading("4. API Usage", level=1)
doc.add_paragraph("POST Request → http://localhost:8000/chain")

doc.add_heading("Request Body:", level=2)
doc.add_paragraph(
    '{\n  "input": {\n    "language": "French",\n    "text": "Hello, how are you?"\n  }\n}'
)

doc.add_heading("Response:", level=2)
doc.add_paragraph(
    '{\n  "output": "Bonjour, comment allez-vous ?"\n}'
)

# Section 5
doc.add_heading("5. Interactive API Docs", level=1)
doc.add_paragraph(
    "Swagger UI: http://localhost:8000/docs\n"
    "ReDoc: http://localhost:8000/redoc"
)

# Section 6
doc.add_heading("6. Directory Structure", level=1)
doc.add_paragraph(
    "project/\n"
    "│── main.py          # FastAPI app\n"
    "│── .env             # API key storage\n"
    "│── requirements.txt # dependencies"
)

# Save document
doc.save("FastAPI_LangChain_Groq_Docs.docx")

print("✅ Documentation saved as FastAPI_LangChain_Groq_Docs.docx")
